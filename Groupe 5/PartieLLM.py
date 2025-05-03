import requests
from docx import Document

# --- 1. Fonctions pour générer les prompts ---

def formatter_formation(formation):
    return "\n".join(
        f"- {f['titre']} – {f['etablissement']} ({f['periode']})\n  " +
        "\n  ".join(f['details']) for f in formation
    )

def formatter_experience(experience):
    return "\n".join(
        f"- {e['titre']} – {e['entreprise']}, {e['lieu']} ({e['periode']})\n  " +
        "\n  ".join(e['details']) for e in experience
    )

def generer_prompt_pertinence(cv_dict, offre_dict):
    formations = formatter_formation(cv_dict["formation"])
    experiences = formatter_experience(cv_dict["experience"])
    competences = "\n- ".join(cv_dict["competences_techniques"])
    soft_skills = "\n- ".join(cv_dict["soft_skills"])
    langues = "\n- ".join(cv_dict["langues"])
    certifications = "\n- ".join(cv_dict["certifications"])

    cv_txt = f"""Nom : {cv_dict['prenom_nom']}
Email : {cv_dict['email']}
Téléphone : {cv_dict['telephone']}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques :
- {competences}

Soft skills :
- {soft_skills}

Langues :
- {langues}

Certifications :
- {certifications}
"""

    offre_txt = f"""Titre : {offre_dict['titre']}
Entreprise : {offre_dict['entreprise']}
Lieu : {offre_dict['lieu']}
Contrat : {offre_dict['type_contrat']}

Description de l’entreprise :
{offre_dict['description_entreprise']}

Missions :
{offre_dict['missions']}

Profil recherché :
{offre_dict['profil_recherche']}
"""

    return f"""
Tu es un expert RH.

Voici un CV et une offre d'emploi. Réponds uniquement par "oui" si le profil correspond à plus de 70 % à l'offre, sinon réponds "non". Ne donne aucune explication.

--- CV ---
{cv_txt}

--- Offre ---
{offre_txt}
"""

def generer_prompt_lettre(cv_dict, offre_dict):
    formations = formatter_formation(cv_dict["formation"])
    experiences = formatter_experience(cv_dict["experience"])
    competences = ", ".join(cv_dict["competences_techniques"])
    soft_skills = ", ".join(cv_dict["soft_skills"])
    langues = ", ".join(cv_dict["langues"])
    certifications = ", ".join(cv_dict["certifications"])

    return f"""
Tu es un expert RH et spécialiste de la rédaction de lettres de motivation professionnelles. Rédige une lettre complète, prête à être envoyée, en t’appuyant sur le CV du candidat et l’offre d’emploi ci-dessous.

🎯 Objectif :
Fournir une lettre claire, convaincante, personnalisée, sans faute ni besoin de correction, dans un style fluide, professionnel et humain.

✅ La lettre doit impérativement :
- Tenir sur une page (Word A4) avec un style direct et efficace.
- Suivre ce plan structuré :
    1. Présentation brève du candidat et de son parcours
    2. Motivation sincère et cohérente pour le poste
    3. Mise en lien entre l’entreprise/l’offre et les valeurs du candidat
    4. Mise en avant ciblée des compétences, expériences ou cours suivis correspondant aux missions
    5. Remerciements, disponibilité pour un entretien, et formule de politesse

✍️ Style :
- Zéro faute d’orthographe ou de grammaire.
- Chaque phrase commence par une majuscule.
- Aucune formule générique ni tournure artificielle.
- Le ton doit être confiant, positif, professionnel et chaleureux.
- Ne propose aucun espace à compléter : tout doit être finalisé.

📎 Contexte fourni :

--- CV ---
Nom : {cv_dict['prenom_nom']}
Email : {cv_dict['email']}
Téléphone : {cv_dict['telephone']}
LinkedIn : {cv_dict['linkedin']}
GitHub : {cv_dict['github']}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques : {competences}
Compétences comportementales (soft skills) : {soft_skills}
Langues : {langues}
Certifications : {certifications}

--- Offre ---
Titre : {offre_dict['titre']}
Entreprise : {offre_dict['entreprise']}
Lieu : {offre_dict['lieu']}
Type de contrat : {offre_dict['type_contrat']}
À propos de l’entreprise :
{offre_dict['description_entreprise']}

Missions proposées :
{offre_dict['missions']}

Profil recherché :
{offre_dict['profil_recherche']}
"""

# --- 2. CV d'exemple ---
cv_dict = {
    "prenom_nom": "Eleonore VERNE",
    "email": "eleonore.verne@marine-analytics.com",
    "telephone": "+377 98 76 54 32",
    "linkedin": "eleonore-verne",
    "github": "eleonoreverne",
    "competences_techniques": ["Python", "R", "Tableau", "SQL", "MATLAB", "Azure"],
    "soft_skills": ["Leadership", "Innovation", "Résolution de problèmes", "Communication"],
    "langues": ["Français (Natif)", "Anglais (C1)", "Espagnol (B2)", "Japonais (B1)"],
    "certifications": ["Yacht Master", "Data Science Professional (DSP-M278X93)"],
    "formation": [
        {
            "titre": "Master Intelligence Artificielle et Politiques Publiques",
            "etablissement": "Sciences Po Paris",
            "periode": "Sept. 2023 – Juin 2025",
            "details": [
                "Principaux enseignements: Deep Learning, Gouvernance des données, Éthique de l’IA, Politiques environnementales, Modélisation prédictive."
            ]
        },
        {
            "titre": "Licence Mathématiques Appliquées et Sciences Sociales",
            "etablissement": "Université Côte d’Azur",
            "periode": "Sept. 2020 – Juin 2023",
            "details": [
                "Principaux enseignements: Statistiques avancées, Économétrie, Analyse de données, Optimisation, Simulation stochastique, Modélisation.",
                "Mention: Très Bien"
            ]
        },
        {
            "titre": "Baccalauréat Scientifique International",
            "etablissement": "Lycée Albert Premier",
            "periode": "Sept. 2017 – Juillet 2020",
            "details": [
                "Spécialité: Mathématiques",
                "Option: Section Européenne Anglais",
                "Mention: Très Bien avec Félicitations du Jury"
            ]
        }
    ],
    "experience": [
        {
            "titre": "Data Scientist Junior",
            "entreprise": "Marine Analytics Monaco",
            "lieu": "Monaco",
            "periode": "Juin 2022 – Déc. 2022",
            "details": [
                "Conception et implémentation d’algorithmes d’analyse prédictive pour optimiser les trajets maritimes et réduire l’empreinte carbone des navires de luxe. Développement d’un tableau de bord interactif permettant le suivi en temps réel des performances et des économies de carburant."
            ]
        },
        {
            "titre": "Stage en Data Analytics",
            "entreprise": "Azur Innovations",
            "lieu": "Nice, France",
            "periode": "Mai 2021 – Août 2021",
            "details": [
                "Analyse des données touristiques de la Côte d’Azur pour identifier les tendances saisonnières et développer un modèle de prévision de l’affluence. Création de visualisations interactives et participation à l’élaboration d’une stratégie de tourisme durable pour les municipalités partenaires."
            ]
        }
    ]
}

# --- 3. Offres d'alternance exemple ---
liste_offres = [
    {
        "titre": "Alternante Data Analyst Junior",
        "entreprise": "NavOcean",
        "lieu": "Marseille",
        "type_contrat": "Alternance (12 mois)",
        "description_entreprise": "NavOcean développe des solutions IA pour la logistique maritime durable.",
        "missions": "- Analyse des données environnementales maritimes\n- Création de tableaux de bord avec Tableau/Power BI\n- Contribution aux modèles prédictifs avec Python/R",
        "profil_recherche": "- Étudiant(e) en Master IA ou Data Science\n- Bon niveau en Python ou R, visualisation et statistiques\n- Sensibilité aux enjeux maritimes ou climatiques"
    },
    {
        "titre": "Alternante Data Scientist Climat",
        "entreprise": "GreenForecast",
        "lieu": "Lyon",
        "type_contrat": "Alternance (12 mois)",
        "description_entreprise": "GreenForecast prédit les impacts climatiques pour les grandes villes européennes.",
        "missions": "- Traitement de données satellites\n- Modélisation prédictive avec Python\n- Veille scientifique sur l'adaptation climatique",
        "profil_recherche": "- Solides compétences Python, Machine Learning, climat\n- Sensibilité environnementale"
    }
]

# --- 4. Envoi à l'API Gemini ---
api_key = "Votre clé API"  # Remplace par ta vraie clé
url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro:generateContent"
headers = {"Content-Type": "application/json"}
params = {"key": api_key}

def interroger_gemini(prompt):
    data = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": prompt}]
            }
        ]
    }
    response = requests.post(url, headers=headers, params=params, json=data)
    if response.status_code == 200:
        resultat = response.json()
        return resultat['candidates'][0]['content']['parts'][0]['text'].strip()
    else:
        print(f"\n❌ Erreur Gemini : {response.status_code}")
        print(response.text)
        return None

# --- 5. Traitement ---
for offre in liste_offres:
    print(f"\n🔍 Traitement de l'offre chez {offre['entreprise']}...")
    prompt_pertinence = generer_prompt_pertinence(cv_dict, offre)
    reponse = interroger_gemini(prompt_pertinence)

    if reponse and reponse.lower() == "oui":
        print("✅ Profil pertinent. Génération de la lettre...")
        prompt_lettre = generer_prompt_lettre(cv_dict, offre)
        lettre = interroger_gemini(prompt_lettre)

        if lettre:
            doc = Document()
            doc.add_heading(f"Lettre de motivation – {cv_dict['prenom_nom']}", 0)
            for ligne in lettre.split('\n'):
                if ligne.strip():
                    doc.add_paragraph(ligne)

            nom_fichier = f"Lettre_{cv_dict['prenom_nom'].replace(' ', '_')}_{offre['entreprise'].replace(' ', '_')}.docx"
            doc.save(nom_fichier)
            print(f"📄 Lettre sauvegardée dans : {nom_fichier}")
        else:
            print("⚠️ Erreur lors de la génération de la lettre.")
    else:
        print("⛔ Profil non jugé pertinent.")

