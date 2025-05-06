import discord
from discord import app_commands
import requests
import json
import io
from docx import Document
from utils.helper import get_user_data, check_user_prerequisites, UserData

# --- 1. Fonctions pour générer les prompts ---

def formatter_formation(formation):
    """Formate une liste de formations pour un prompt"""
    if not formation:
        return "Aucune formation mentionnée"
    
    return "\n".join(
        f"- {f['titre']} – {f.get('etablissement', 'N/A')} ({f.get('periode', 'N/A')})\n  " +
        "\n  ".join(f.get('details', [])) for f in formation
    )

def formatter_experience(experience):
    """Formate une liste d'expériences pour un prompt"""
    if not experience:
        return "Aucune expérience professionnelle mentionnée"
    
    return "\n".join(
        f"- {e['titre']} – {e.get('entreprise', 'N/A')}, {e.get('lieu', 'N/A')} ({e.get('periode', 'N/A')})\n  " +
        "\n  ".join(e.get('details', [])) for e in experience
    )

def generer_prompt_pertinence(cv_dict, offre_dict):
    """Génère un prompt pour vérifier la pertinence du CV pour l'offre"""
    formations = formatter_formation(cv_dict.get("formation", []))
    experiences = formatter_experience(cv_dict.get("experience", []))
    competences = "\n- ".join(cv_dict.get("competences_techniques", ["Aucune compétence technique mentionnée"]))
    soft_skills = "\n- ".join(cv_dict.get("soft_skills", ["Aucun soft skill mentionné"]))
    langues = "\n- ".join(cv_dict.get("langues", ["Aucune langue mentionnée"]))
    certifications = "\n- ".join(cv_dict.get("certifications", ["Aucune certification mentionnée"]))

    cv_txt = f"""Nom : {cv_dict.get('prenom_nom', 'Non spécifié')}
Email : {cv_dict.get('email', 'Non spécifié')}
Téléphone : {cv_dict.get('telephone', 'Non spécifié')}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques :
- {competences}

Soft skills :
- {soft_skills}

Langues :
- {langues}

Certifications :
- {certifications}
"""

    offre_txt = f"""Titre : {offre_dict.get('titre', 'Non spécifié')}
Entreprise : {offre_dict.get('entreprise', 'Non spécifié')}
Lieu : {offre_dict.get('lieu', 'Non spécifié')}
Contrat : {offre_dict.get('type_contrat', 'Non spécifié')}

Description de l'entreprise :
{offre_dict.get('description_entreprise', 'Non spécifié')}

Missions :
{offre_dict.get('missions', 'Non spécifié')}

Profil recherché :
{offre_dict.get('profil_recherche', 'Non spécifié')}
"""

    return f"""
Tu es un expert RH.

Voici un CV et une offre d'emploi. Analyse en détail la compatibilité entre le profil et le poste.
1. Réponds d'abord par "oui" si le profil correspond à plus de 70 % à l'offre, sinon réponds "non".
2. Ensuite, liste les principales forces du candidat par rapport au poste (3 à 5 points).
3. Liste les éventuels points à améliorer ou compétences manquantes (2 à 3 points).
4. Donne un pourcentage approximatif de correspondance.

--- CV ---
{cv_txt}

--- Offre ---
{offre_txt}
"""

def generer_prompt_lettre(cv_dict, offre_dict, infos_perso=None):
    """Génère un prompt pour créer une lettre de motivation"""
    if infos_perso is None:
        infos_perso = {"motivation": "", "lien_entreprise": "", "contraintes": ""}

    formations = formatter_formation(cv_dict.get("formation", []))
    experiences = formatter_experience(cv_dict.get("experience", []))
    competences = ", ".join(cv_dict.get("competences_techniques", ["Aucune compétence technique mentionnée"]))
    soft_skills = ", ".join(cv_dict.get("soft_skills", ["Aucun soft skill mentionné"]))
    langues = ", ".join(cv_dict.get("langues", ["Aucune langue mentionnée"]))
    certifications = ", ".join(cv_dict.get("certifications", ["Aucune certification mentionnée"]))

    perso_txt = ""
    if infos_perso["motivation"]:
        perso_txt += f"\nMotivation personnelle : {infos_perso['motivation']}"
    if infos_perso["lien_entreprise"]:
        perso_txt += f"\nLien particulier avec l'entreprise ou le secteur : {infos_perso['lien_entreprise']}"
    if infos_perso["contraintes"]:
        perso_txt += f"\nInformations supplémentaires : {infos_perso['contraintes']}"

    return f"""
Tu es un expert RH et spécialiste de la rédaction de lettres de motivation professionnelles. Rédige une lettre complète, prête à être envoyée, en t'appuyant sur le CV du candidat et l'offre d'emploi ci-dessous.

🎯 Objectif :
Fournir une lettre claire, convaincante, personnalisée, sans faute ni besoin de correction, dans un style fluide, professionnel et humain.

✅ La lettre doit impérativement :
- Tenir sur une page (Word A4) avec un style direct et efficace.
- Suivre ce plan structuré :
    1. Présentation brève du candidat et de son parcours
    2. Motivation sincère et cohérente pour le poste
    3. Mise en lien entre l'entreprise/l'offre et les valeurs du candidat
    4. Mise en avant ciblée des compétences, expériences ou cours suivis correspondant aux missions
    5. Remerciements, disponibilité pour un entretien, et formule de politesse

✍️ Style :
- Zéro faute d'orthographe ou de grammaire.
- Chaque phrase commence par une majuscule.
- Aucune formule générique ni tournure artificielle.
- Le ton doit être confiant, positif, professionnel et chaleureux.
- Ne propose aucun espace à compléter : tout doit être finalisé.

📎 Contexte fourni :

--- CV ---
Nom : {cv_dict.get('prenom_nom', 'Non spécifié')}
Email : {cv_dict.get('email', 'Non spécifié')}
Téléphone : {cv_dict.get('telephone', 'Non spécifié')}
LinkedIn : {cv_dict.get('linkedin', 'Non spécifié')}
GitHub : {cv_dict.get('github', 'Non spécifié')}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques : {competences}
Compétences comportementales (soft skills) : {soft_skills}
Langues : {langues}
Certifications : {certifications}

--- Offre ---
Titre : {offre_dict.get('titre', 'Non spécifié')}
Entreprise : {offre_dict.get('entreprise', 'Non spécifié')}
Lieu : {offre_dict.get('lieu', 'Non spécifié')}
Type de contrat : {offre_dict.get('type_contrat', 'Non spécifié')}
À propos de l'entreprise :
{offre_dict.get('description_entreprise', 'Non spécifié')}

Missions proposées :
{offre_dict.get('missions', 'Non spécifié')}

Profil recherché :
{offre_dict.get('profil_recherche', 'Non spécifié')}

--- Informations complémentaires du candidat ---
{perso_txt if perso_txt else "Aucune information supplémentaire fournie."}
"""

# --- 2. Interaction avec l'API Gemini ---
async def interroger_gemini(prompt, api_key):
    """Interroge l'API Gemini avec un prompt"""
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro:generateContent"
    headers = {"Content-Type": "application/json"}
    params = {"key": api_key}
    
    data = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": prompt}]
            }
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, params=params, json=data)
        if response.status_code == 200:
            resultat = response.json()
            return resultat['candidates'][0]['content']['parts'][0]['text'].strip()
        else:
            print(f"\n❌ Erreur Gemini : {response.status_code}")
            print(response.text)
            return None
    except Exception as e:
        print(f"Erreur lors de l'appel à l'API Gemini: {e}")
        return None

# --- 3. Fonctions pour le bot Discord ---

def setup_gemini_commands(bot, api_key):
    """Configure les commandes liées à PartieLLM pour le bot Discord"""
    
    # Commande pour comparer le CV avec une offre
    @bot.tree.command(name="analyser_cv_offre", description="Analyse la compatibilité entre votre CV et l'offre d'emploi sélectionnée")
    async def analyser_cv_offre(interaction: discord.Interaction):
        await interaction.response.defer(thinking=True)
        
        # Vérifier que l'utilisateur a téléchargé un CV et sélectionné une offre
        error_message = check_user_prerequisites(interaction.user.id, need_cv=True, need_job_offer=True)
        if error_message:
            await interaction.followup.send(error_message, ephemeral=True)
            return
        
        try:
            user_data = get_user_data(interaction.user.id)
            
            # Vérifier que l'analyse structurée du CV est disponible
            if not isinstance(user_data.cv_structured, dict):
                await interaction.followup.send("❌ Votre CV n'a pas été analysé de manière structurée. Utilisez d'abord la commande `/extraire_cv` pour l'analyser.", ephemeral=True)
                return
                
            # Générer le prompt pour l'analyse de pertinence
            prompt = generer_prompt_pertinence(user_data.cv_structured, user_data.job_offer)
            
            # Interroger Gemini
            response = await interroger_gemini(prompt, api_key)
            
            if not response:
                await interaction.followup.send("❌ Une erreur s'est produite lors de l'analyse avec Gemini.", ephemeral=True)
                return
                
            # Analyse de la réponse pour en tirer les informations clés
            lines = response.split('\n')
            decision = "Non déterminé"
            matching_percentage = "N/A"
            strengths = []
            improvements = []
            
            for i, line in enumerate(lines):
                if i == 0 and ("oui" in line.lower() or "non" in line.lower()):
                    decision = line.strip()
                elif "%" in line:
                    matching_percentage = line.strip()
                elif line.startswith("-") or line.startswith("•") or (i > 0 and any(s in line.lower() for s in ["force", "point fort", "avantage"])):
                    strengths.append(line.strip())
                elif any(s in line.lower() for s in ["amélioration", "manquant", "lacune", "faiblesse"]):
                    improvements.append(line.strip())
            
            # Créer un embed Discord pour présenter les résultats
            is_match = "oui" in decision.lower()
            color = discord.Color.green() if is_match else discord.Color.red()
            
            embed = discord.Embed(
                title=f"Analyse de compatibilité: {user_data.job_offer['titre']}",
                description=f"**Décision**: {decision}\n**Correspondance**: {matching_percentage}",
                color=color
            )
            
            if strengths:
                embed.add_field(name="📈 Points forts", value="\n".join([f"✅ {s}" for s in strengths[:5]]), inline=False)
            
            if improvements:
                embed.add_field(name="🔍 Points à améliorer", value="\n".join([f"❗ {i}" for i in improvements[:3]]), inline=False)
            
            # Encourager à générer une lettre si le match est bon
            if is_match:
                embed.add_field(
                    name="💡 Prochaine étape", 
                    value="Votre profil correspond bien à cette offre! Utilisez `/infos_lettre_g5` pour générer une lettre de motivation personnalisée.", 
                    inline=False
                )
            
            await interaction.followup.send(embed=embed)
            
        except Exception as e:
            print(f"Erreur lors de l'analyse CV/offre: {e}")
            await interaction.followup.send(f"❌ Une erreur s'est produite: {str(e)}", ephemeral=True)
    
    # Commande pour collecter des informations supplémentaires pour la lettre
    @bot.tree.command(name="infos_lettre_g5", description="Fournir des informations supplémentaires pour personnaliser votre lettre de motivation")
    async def infos_lettre_g5(interaction: discord.Interaction):
        # Vérifier que l'utilisateur a téléchargé un CV et sélectionné une offre
        error_message = check_user_prerequisites(interaction.user.id, need_cv=True, need_job_offer=True)
        if error_message:
            await interaction.response.send_message(error_message, ephemeral=True)
            return
        
        # Créer un modal pour collecter les informations
        class InfosLettreModal(discord.ui.Modal, title="Informations pour votre lettre de motivation"):
            motivation = discord.ui.TextInput(
                label="Votre motivation personnelle pour ce poste",
                style=discord.TextStyle.paragraph,
                placeholder="Pourquoi ce poste vous intéresse particulièrement...",
                required=False,
                max_length=1000
            )
            
            lien_entreprise = discord.ui.TextInput(
                label="Lien avec l'entreprise ou le secteur",
                style=discord.TextStyle.paragraph,
                placeholder="Expérience, intérêt ou connaissance particulière...",
                required=False,
                max_length=1000
            )
            
            contraintes = discord.ui.TextInput(
                label="Informations supplémentaires",
                style=discord.TextStyle.paragraph,
                placeholder="Contraintes géographiques, disponibilité, etc.",
                required=False,
                max_length=1000
            )
            
            async def on_submit(self, interaction: discord.Interaction):
                user_data = get_user_data(interaction.user.id)
                
                # Stocker les informations dans les données utilisateur
                if not hasattr(user_data, "lettre_infos"):
                    user_data.lettre_infos = {}
                
                user_data.lettre_infos = {
                    "motivation": self.motivation.value,
                    "lien_entreprise": self.lien_entreprise.value,
                    "contraintes": self.contraintes.value
                }
                
                # Confirmer et suggérer la prochaine étape
                await interaction.response.send_message(
                    "✅ Vos informations ont bien été enregistrées. Utilisez `/generer_lettre_g5` pour créer votre lettre de motivation personnalisée.",
                    ephemeral=True
                )
        
        # Montrer le modal
        await interaction.response.send_modal(InfosLettreModal())
    
    # Commande pour générer la lettre de motivation
    @bot.tree.command(name="generer_lettre_g5", description="Générer une lettre de motivation basée sur votre CV et l'offre d'emploi avec Gemini")
    async def generer_lettre_g5(interaction: discord.Interaction):
        await interaction.response.defer(thinking=True)
        
        # Vérifier que l'utilisateur a téléchargé un CV et sélectionné une offre
        error_message = check_user_prerequisites(interaction.user.id, need_cv=True, need_job_offer=True)
        if error_message:
            await interaction.followup.send(error_message, ephemeral=True)
            return
        
        try:
            user_data = get_user_data(interaction.user.id)
            
            # Vérifier que l'analyse structurée du CV est disponible
            if not isinstance(user_data.cv_structured, dict):
                await interaction.followup.send("❌ Votre CV n'a pas été analysé de manière structurée. Utilisez d'abord la commande `/extraire_cv` pour l'analyser.", ephemeral=True)
                return
            
            # Récupérer les infos supplémentaires si elles existent
            infos_perso = getattr(user_data, "lettre_infos", {"motivation": "", "lien_entreprise": "", "contraintes": ""})
            
            # Générer le prompt pour la lettre
            prompt = generer_prompt_lettre(user_data.cv_structured, user_data.job_offer, infos_perso)
            
            # Interroger Gemini
            lettre = await interroger_gemini(prompt, api_key)
            
            if not lettre:
                await interaction.followup.send("❌ Une erreur s'est produite lors de la génération de la lettre avec Gemini.", ephemeral=True)
                return
            
            # Créer un document Word avec la lettre
            doc = Document()
            
            # Ajouter un en-tête
            nom = user_data.cv_structured.get("prenom_nom", "")
            doc.add_heading(f"Lettre de motivation - {nom}", 0)
            
            # Ajouter le contenu de la lettre
            for ligne in lettre.split('\n'):
                if ligne.strip():
                    doc.add_paragraph(ligne)
            
            # Sauvegarder en mémoire
            with io.BytesIO() as file:
                doc.save(file)
                file.seek(0)
                
                # Préparer le fichier pour Discord
                entreprise = user_data.job_offer.get("entreprise", "entreprise").replace(" ", "_")
                file_discord = discord.File(file, filename=f"Lettre_Motivation_{nom.replace(' ', '_')}_{entreprise}.docx")
                
                # Créer un embed pour la présentation
                embed = discord.Embed(
                    title=f"📝 Lettre de motivation pour {user_data.job_offer.get('titre', 'le poste')}",
                    description=f"Votre lettre de motivation personnalisée pour {user_data.job_offer.get('entreprise', 'l\'entreprise')} est prête!",
                    color=discord.Color.blue()
                )
                
                embed.add_field(
                    name="💡 Conseil",
                    value="N'oubliez pas de relire et personnaliser davantage cette lettre avant de l'envoyer.",
                    inline=False
                )
                
                # Envoyer le fichier
                await interaction.followup.send(embed=embed, file=file_discord)
            
        except Exception as e:
            print(f"Erreur lors de la génération de la lettre: {e}")
            await interaction.followup.send(f"❌ Une erreur s'est produite: {str(e)}", ephemeral=True)

# Fonction d'initialisation pour être importée dans bot.py
def setup_partillm_commands(bot, api_key):
    setup_gemini_commands(bot, api_key)